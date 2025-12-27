import streamlit as st
from docx import Document
import os

st.set_page_config(page_title="Form A Generator")

st.title("Form A DOCX Generator")

if st.button("Generate Form A"):
    doc = Document()
    doc.add_heading("FORM A", level=1)
    doc.add_paragraph("Mediation Application")

    file_name = "FormA.docx"
    doc.save(file_name)

    with open(file_name, "rb") as f:
        st.download_button(
            label="Download Form A",
            data=f,
            file_name=file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
