from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from docx.enum.table import WD_TABLE_ALIGNMENT

doc= Document()


def add_top_text(doc, text, size=14, bold=False):
    p= doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.4
    
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


add_top_text(doc, "FORM 'A'", size=14, bold=True)
add_top_text(doc, "MEDIATION APPLICATION FOR", size=14, bold=True)
add_top_text(doc, "[REFER RULE 3(1)]", size=12, bold=True)
add_top_text(doc, "Mumbai District Legal Services Authority", size=11)
add_top_text(doc, "City Civil Court Mumbai", size=11)



table = doc.add_table(rows=1, cols=2)

def cell_text(cell, text,bold=False, size=11):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing=1
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

def telephone_mobile_email(tele_no,email_id=None):
    if email_id is not None:
        row= table.add_row()
        row.cells[0].text =""
        cell_text(row.cells[1], "Telephone No.", bold=True)
        cell_text(row.cells[2], "{{mobile}}", bold=True)


        row= table.add_row()
        row.cells[0].text =""
        cell_text(row.cells[1], "Mobile Number", bold=True)
        cell_text(row.cells[2], "", bold=True)
        
        row = table.add_row()
        row.cells[0].text =""
        cell_text(row.cells[1],"Email ID", bold=True)
        
        cell = row.cells[2]
        cell.text = ""

        p=cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("info@kslegal.co.in")
        run.underline = True

    elif email_id is None:
        row= table.add_row()
        row.cells[0].text =""
        cell_text(row.cells[1], "Telephone no.", bold=True)
        row.cells[2].text =""


        row= table.add_row()
        row.cells[0].text =""
        cell_text(row.cells[1], "Mobile Number", bold=True)
        cell_text(row.cells[2], "", bold=True)
        
        row = table.add_row()
        row.cells[0].text =""
        cell_text(row.cells[1],"Email ID", bold=True)

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
table.autofit = False

table.columns[0].width = Inches(0.2)
table.columns[1].width = Inches(2)
table.columns[2].width = Inches(2.5)



row= table.rows[0]
row.cells[0].merge(row.cells[2])
cell_text(row.cells[0], "DETAIL OF PARTIES:", bold=True)


row = table.add_row()
cell_text(row.cells[0],"1")
cell_text(row.cells[1], "Name of Applicant", bold=True)
cell_text(row.cells[2],"{{client_name}}", bold=True)

row = table.add_row()
row.cells[0].text = ""
row.cells[1].merge(row.cells[2])
cell_text(row.cells[1],"Address and contact details of Applicant", bold = True)


row = table.add_row()
cell_text(row.cells[0], "1", bold=True)
cell_text(row.cells[1], "Adddress", bold=True)
cell_text(row.cells[2], "Registered Address\n {{branch_address}}\n\n\n"
                         
                         "Correspondance Branch Address\n {{branch_address}}\n", bold=True)


telephone_mobile_email(tele_no="Telephone", email_id="info@kslegal.co.in")


row = table.add_row()
row.cells[0].text = "2"
row.cells[1].merge(row.cells[2])
cell_text(row.cells[1], "Name, Address and Contact details of Opposite Party:", bold=True)

row = table.add_row()
row.cells[0].text=""
row.cells[1].merge(row.cells[2])
cell_text(row.cells[1],"Address and contact details of Defendant/s", bold=True, size=10)


row = table.add_row()
row.cells[0].text=""
cell_text(row.cells[1],"Name", bold=True)
cell_text(row.cells[2],"{{customer_name}}")


row = table.add_row()
row.cells[0].text=""
cell_text(row.cells[1],"Address", bold=True)
cell_text(row.cells[2],"REGISTERED ADDRESS:\n {% if address1 and address1 != "" %}{{address1}} {% else %} ________________ {%endif %} \n\n"
           "CORRESPONDENCE ADDRESS:\n {% if address1 and address1 != "" %}{{address1}} {% else %} ________________ {%endif %}\n")



telephone_mobile_email(tele_no="", email_id=None)



row=table.add_row()
row.cells[0].merge(row.cells[2])
cell_text(row.cells[0],"DETAILS OF DISPUTE:", bold=True)


row = table.add_row()
row.cells[0].merge(row.cells[2])
cell = row.cells[0]  
p = cell.paragraphs[0]

p.alignment = WD_ALIGN_PARAGRAPH.CENTER

run = p.add_run("THE COMM. COURTS (PRE-INSTITUTION....SETTLEMENT) RULES, 2018")
run.bold = True
run.underline = True


row = table.add_row()
row.cells[0].text =""
row.cells[1].merge(row.cells[2])
cell_text(row.cells[1],"Nature of disputes as per section 2(1)(c) of the Commercial Courts Act, 2015 (4 of 2016):", bold=True, size=10)



        


doc.save("FormA.docx")


